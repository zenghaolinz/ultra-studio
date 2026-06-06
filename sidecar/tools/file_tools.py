import json
import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".json",
    ".jsonl",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".css",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".py",
    ".rs",
    ".toml",
    ".ini",
    ".log",
}

READABLE_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx"}
MAX_READ_BYTES = 2_000_000
MAX_LIST_ITEMS = 300
MAX_MANY_FILES = 30
MAX_SEARCH_MATCHES = 200
MAX_COMMAND_OUTPUT = 24000
DANGEROUS_COMMAND_PATTERNS = [
    r"\bdel\s+/(?:s|q|f)",
    r"\brmdir\s+/(?:s|q)",
    r"\brd\s+/(?:s|q)",
    r"\bformat\b",
    r"\bdiskpart\b",
    r"\bRemove-Item\b.*(?<!\w)-Recurse\b",
    r"\brm\s+-rf\b",
    r"\bgit\s+reset\s+--hard\b",
    r"\bgit\s+clean\s+-fd",
]


def _norm(path: str) -> Path:
    expanded = os.path.expandvars(os.path.expanduser(path.strip()))
    lowered = expanded.lower().replace("\\", "/")
    if lowered in {"desktop", "桌面"}:
        expanded = str(Path.home() / "Desktop")
    elif lowered.startswith("desktop/") or lowered.startswith("桌面/"):
        _, rest = expanded.replace("\\", "/", 1).split("/", 1)
        expanded = str(Path.home() / "Desktop" / rest)
    return Path(expanded).resolve()


def _ensure_docx_path(file_path: str) -> Path:
    path = _norm(file_path)
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    return path


def _truncate(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars] + "\n\n[内容已截断]", True


def _read_text_file(path: Path) -> str:
    data = path.read_bytes()
    if len(data) > MAX_READ_BYTES:
        data = data[:MAX_READ_BYTES]
    for enc in ("utf-8-sig", "utf-8", "gbk", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("缺少 pypdf 依赖，请先安装 requirements.txt") from exc

    reader = PdfReader(str(path))
    parts = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        parts.append(f"\n--- Page {index} ---\n{text}")
    return "\n".join(parts).strip()


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("缺少 python-docx 依赖，请先安装 requirements.txt") from exc

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _load_docx_document(path: Path):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("Missing python-docx dependency. Please install sidecar requirements.txt") from exc
    return Document(str(path))


def create_docx_document(
    file_path: str,
    title: str = "",
    paragraphs: list[str] | None = None,
    overwrite: bool = False,
) -> dict[str, Any]:
    path = _ensure_docx_path(file_path)
    if path.exists() and not overwrite:
        path = _unique_destination(path)

    try:
        from docx import Document
    except Exception as exc:
        return {
            "ok": False,
            "error": "Missing python-docx dependency. Please install sidecar requirements.txt",
        }

    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        if title.strip():
            doc.add_heading(title.strip(), level=1)
        for paragraph in paragraphs or []:
            doc.add_paragraph(str(paragraph))
        doc.save(str(path))
    except Exception as exc:
        return {"ok": False, "error": str(exc), "path": str(path)}

    return {
        "ok": True,
        "path": str(path),
        "name": path.name,
        "created": True,
        "overwritten": overwrite,
        "paragraph_count": len(paragraphs or []),
    }


def edit_docx_document(
    file_path: str,
    action: str,
    text: str = "",
    find: str = "",
    replace: str = "",
    backup: bool = False,
) -> dict[str, Any]:
    path = _ensure_docx_path(file_path)
    if not path.exists() or not path.is_file():
        return {"ok": False, "error": f"DOCX file not found: {path}"}
    if action not in {"append", "prepend", "replace"}:
        return {"ok": False, "error": "action must be append, prepend, or replace"}

    backup_path = None
    if backup:
        backup_path = _unique_destination(path.with_suffix(path.suffix + ".bak"))
        shutil.copy2(path, backup_path)

    try:
        doc = _load_docx_document(path)
        replacements = 0
        if action == "append":
            doc.add_paragraph(text)
        elif action == "prepend":
            existing = [p.text for p in doc.paragraphs]
            for p in list(doc.paragraphs):
                element = p._element
                element.getparent().remove(element)
            doc.add_paragraph(text)
            for paragraph in existing:
                doc.add_paragraph(paragraph)
        elif action == "replace":
            if not find:
                return {"ok": False, "error": "replace action requires find"}
            for paragraph in doc.paragraphs:
                if find in paragraph.text:
                    paragraph.text = paragraph.text.replace(find, replace)
                    replacements += 1
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if find in cell.text:
                            cell.text = cell.text.replace(find, replace)
                            replacements += 1
            if replacements == 0:
                return {"ok": False, "error": "No matching text found", "path": str(path)}

        doc.save(str(path))
    except Exception as exc:
        return {"ok": False, "error": str(exc), "path": str(path)}

    return {
        "ok": True,
        "path": str(path),
        "changed": True,
        "action": action,
        "replacements": replacements if action == "replace" else 0,
        "backup_path": str(backup_path) if backup_path else None,
    }


def read_document(file_path: str, max_chars: int = 12000) -> dict[str, Any]:
    path = _norm(file_path)
    if not path.exists():
        return {"ok": False, "error": f"文件不存在: {path}"}
    if not path.is_file():
        return {"ok": False, "error": f"不是文件: {path}"}

    ext = path.suffix.lower()
    if ext not in READABLE_EXTENSIONS:
        return {
            "ok": False,
            "error": f"暂不支持读取 {ext or '无扩展名'} 文件。支持: {', '.join(sorted(READABLE_EXTENSIONS))}",
        }

    try:
        if ext == ".pdf":
            text = _read_pdf(path)
        elif ext == ".docx":
            text = _read_docx(path)
        else:
            text = _read_text_file(path)
    except Exception as exc:
        return {"ok": False, "error": str(exc), "path": str(path)}

    content, truncated = _truncate(text, max_chars)
    return {
        "ok": True,
        "path": str(path),
        "name": path.name,
        "extension": ext,
        "size_bytes": path.stat().st_size,
        "truncated": truncated,
        "content": content,
    }


def read_many_files(
    file_paths: list[str],
    max_chars_per_file: int = 8000,
    max_files: int = 12,
) -> dict[str, Any]:
    max_files = max(1, min(max_files, MAX_MANY_FILES))
    max_chars_per_file = max(200, min(max_chars_per_file, 50000))
    results = []
    for raw_path in file_paths[:max_files]:
        result = read_document(str(raw_path), max_chars_per_file)
        results.append(result)
    return {
        "ok": True,
        "count": len(results),
        "truncated": len(file_paths) > max_files,
        "files": results,
    }


def list_directory(directory_path: str, recursive: bool = False, max_items: int = 120) -> dict[str, Any]:
    root = _norm(directory_path)
    if not root.exists():
        return {"ok": False, "error": f"目录不存在: {root}"}
    if not root.is_dir():
        return {"ok": False, "error": f"不是目录: {root}"}

    max_items = max(1, min(max_items, MAX_LIST_ITEMS))
    iterator = root.rglob("*") if recursive else root.iterdir()
    items = []
    for item in iterator:
        if len(items) >= max_items:
            break
        try:
            stat = item.stat()
        except OSError:
            continue
        items.append(
            {
                "name": item.name,
                "path": str(item),
                "type": "dir" if item.is_dir() else "file",
                "extension": item.suffix.lower() if item.is_file() else "",
                "size_bytes": stat.st_size if item.is_file() else 0,
                "modified": stat.st_mtime,
            }
        )

    return {
        "ok": True,
        "path": str(root),
        "recursive": recursive,
        "count": len(items),
        "items": items,
        "truncated": len(items) >= max_items,
    }


def search_files(
    directory_path: str,
    query: str,
    file_glob: str = "*",
    recursive: bool = True,
    search_content: bool = True,
    max_matches: int = 80,
) -> dict[str, Any]:
    root = _norm(directory_path)
    if not root.exists() or not root.is_dir():
        return {"ok": False, "error": f"目录不存在或不是目录: {root}"}
    query = str(query or "")
    if not query:
        return {"ok": False, "error": "query 不能为空", "path": str(root)}
    max_matches = max(1, min(max_matches, MAX_SEARCH_MATCHES))
    ignored_dirs = {".git", "node_modules", "target", "dist", "build", "__pycache__", ".venv"}
    iterator = root.rglob(file_glob or "*") if recursive else root.glob(file_glob or "*")
    matches = []
    lowered_query = query.lower()
    for path in iterator:
        if len(matches) >= max_matches:
            break
        if any(part in ignored_dirs for part in path.parts):
            continue
        if not path.is_file():
            continue
        name_match = lowered_query in path.name.lower()
        content_matches = []
        if search_content and path.suffix.lower() in TEXT_EXTENSIONS and path.stat().st_size <= MAX_READ_BYTES:
            try:
                for line_no, line in enumerate(_read_text_file(path).splitlines(), start=1):
                    if lowered_query in line.lower():
                        content_matches.append({"line": line_no, "text": line[:300]})
                        if len(content_matches) >= 5:
                            break
            except Exception:
                pass
        if name_match or content_matches:
            matches.append(
                {
                    "path": str(path),
                    "name": path.name,
                    "extension": path.suffix.lower(),
                    "name_match": name_match,
                    "matches": content_matches,
                }
            )
    return {
        "ok": True,
        "path": str(root),
        "query": query,
        "count": len(matches),
        "matches": matches,
        "truncated": len(matches) >= max_matches,
    }


def _category_for(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"}:
        return "Images"
    if ext in {".pdf", ".docx", ".doc", ".txt", ".md", ".xlsx", ".xls", ".csv", ".pptx", ".ppt"}:
        return "Documents"
    if ext in {".glb", ".gltf", ".obj", ".fbx", ".stl", ".blend"}:
        return "3D"
    if ext in {".mp4", ".mov", ".avi", ".mkv", ".wav", ".mp3", ".flac"}:
        return "Media"
    if ext in {".zip", ".rar", ".7z", ".tar", ".gz"}:
        return "Archives"
    return "Other"


def _unique_destination(path: Path) -> Path:
    if not path.exists():
        return path
    stem, suffix = path.stem, path.suffix
    parent = path.parent
    for index in range(1, 1000):
        candidate = parent / f"{stem}_{index}{suffix}"
        if not candidate.exists():
            return candidate
    raise RuntimeError(f"无法生成不冲突的目标文件名: {path}")


def organize_files(
    directory_path: str,
    strategy: str = "by_type",
    apply_changes: bool = False,
    recursive: bool = False,
) -> dict[str, Any]:
    root = _norm(directory_path)
    if not root.exists() or not root.is_dir():
        return {"ok": False, "error": f"目录不存在或不是目录: {root}"}
    if strategy not in {"by_type", "by_extension"}:
        return {"ok": False, "error": "strategy 只支持 by_type 或 by_extension"}

    iterator = root.rglob("*") if recursive else root.iterdir()
    operations = []
    for item in iterator:
        if not item.is_file():
            continue
        if "logs" in item.parts and item.name == "ultra-studio-sidecar.log":
            continue
        folder_name = _category_for(item) if strategy == "by_type" else (item.suffix.lower().lstrip(".") or "no_extension")
        dest_dir = root / folder_name
        if item.parent == dest_dir:
            continue
        dest = _unique_destination(dest_dir / item.name)
        operations.append({"from": str(item), "to": str(dest)})

    if apply_changes:
        for op in operations:
            dest = Path(op["to"])
            dest.parent.mkdir(parents=True, exist_ok=True)
            shutil.move(op["from"], op["to"])

    return {
        "ok": True,
        "path": str(root),
        "strategy": strategy,
        "applied": apply_changes,
        "operation_count": len(operations),
        "operations": operations[:MAX_LIST_ITEMS],
        "truncated": len(operations) > MAX_LIST_ITEMS,
    }


def edit_text_file(
    file_path: str,
    action: str,
    text: str = "",
    find: str = "",
    replace: str = "",
    use_regex: bool = False,
    backup: bool = False,
) -> dict[str, Any]:
    path = _norm(file_path)
    if not path.exists() or not path.is_file():
        return {"ok": False, "error": f"文件不存在或不是文件: {path}"}
    if path.suffix.lower() not in TEXT_EXTENSIONS:
        return {"ok": False, "error": f"为安全起见，只允许修改文本文件: {path.suffix}"}
    if path.stat().st_size > MAX_READ_BYTES:
        return {"ok": False, "error": f"文件过大，拒绝自动修改: {path}"}
    if action not in {"append", "prepend", "replace"}:
        return {"ok": False, "error": "action 只支持 append、prepend、replace"}

    original = _read_text_file(path)
    updated = original
    replacements = 0

    if action == "append":
        updated = original + ("" if original.endswith("\n") else "\n") + text
    elif action == "prepend":
        updated = text + ("" if text.endswith("\n") else "\n") + original
    elif action == "replace":
        if not find:
            return {"ok": False, "error": "replace 操作必须提供 find"}
        if use_regex:
            updated, replacements = re.subn(find, replace, original)
        else:
            replacements = original.count(find)
            updated = original.replace(find, replace)
        if replacements == 0:
            return {"ok": False, "error": "未找到要替换的内容", "path": str(path)}

    if updated == original:
        return {"ok": True, "path": str(path), "changed": False, "message": "内容没有变化"}

    backup_path = None
    if backup:
        backup_path = _unique_destination(path.with_suffix(path.suffix + ".bak"))
        backup_path.write_text(original, encoding="utf-8")

    path.write_text(updated, encoding="utf-8")
    return {
        "ok": True,
        "path": str(path),
        "changed": True,
        "action": action,
        "replacements": replacements,
        "backup_path": str(backup_path) if backup_path else None,
    }


def write_many_files(
    root_path: str,
    files: list[dict[str, Any]],
    overwrite: bool = False,
) -> dict[str, Any]:
    root = _norm(root_path)
    root.mkdir(parents=True, exist_ok=True)
    if not root.is_dir():
        return {"ok": False, "error": f"不是目录: {root}"}
    written = []
    errors = []
    for item in files[:MAX_MANY_FILES]:
        raw_name = str(item.get("path") or item.get("filename") or item.get("name") or "").replace("\\", "/")
        content = item.get("content")
        if not raw_name or content is None:
            errors.append({"path": raw_name, "error": "缺少 path/filename 或 content"})
            continue
        parts = [part for part in raw_name.lstrip("/").split("/") if part not in {"", ".", ".."}]
        if not parts:
            errors.append({"path": raw_name, "error": "无效路径"})
            continue
        target = (root / Path(*parts)).resolve()
        if root not in target.parents and target != root:
            errors.append({"path": raw_name, "error": "路径越界"})
            continue
        if target.suffix.lower() not in TEXT_EXTENSIONS and target.name.lower() not in {"dockerfile", "makefile", "readme", "license", ".gitignore"}:
            target = target.with_suffix(".txt")
        if target.exists() and not overwrite:
            target = _unique_destination(target)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(str(content), encoding="utf-8")
        written.append({"path": str(target), "name": target.name, "size_bytes": target.stat().st_size})
    return {
        "ok": len(written) > 0 and not errors,
        "root": str(root),
        "written_count": len(written),
        "error_count": len(errors),
        "files": written,
        "errors": errors,
        "truncated": len(files) > MAX_MANY_FILES,
    }


def _command_is_dangerous(command: str) -> bool:
    return any(re.search(pattern, command, re.IGNORECASE) for pattern in DANGEROUS_COMMAND_PATTERNS)


def run_command(
    command: str,
    cwd: str = "",
    shell: str = "powershell",
    timeout_seconds: int = 60,
    confirmed: bool = False,
    permission_mode: str = "standard",
) -> dict[str, Any]:
    command = str(command or "").strip()
    if not command:
        return {"ok": False, "error": "command 不能为空"}
    workdir = _norm(cwd) if cwd else Path.cwd()
    if not workdir.exists() or not workdir.is_dir():
        return {"ok": False, "error": f"工作目录不存在或不是目录: {workdir}"}
    if _command_is_dangerous(command):
        return {"ok": False, "error": "命令包含高风险操作，已拒绝执行。", "command": command, "cwd": str(workdir)}
    if permission_mode != "autonomous" and not confirmed:
        return {
            "ok": False,
            "needs_confirmation": True,
            "command": command,
            "cwd": str(workdir),
            "message": "\n".join(
                [
                    "[CONFIRM_COMMAND_REQUIRED]",
                    f"命令: `{command}`",
                    f"目录: `{workdir}`",
                    "提示: 标准模式下执行系统命令前需要确认。",
                    "[/CONFIRM_COMMAND_REQUIRED]",
                ]
            ),
        }
    timeout_seconds = max(1, min(int(timeout_seconds or 60), 300))
    if shell == "cmd":
        args = ["cmd.exe", "/c", command]
    else:
        args = ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", command]
    try:
        completed = subprocess.run(
            args,
            cwd=str(workdir),
            text=True,
            capture_output=True,
            timeout=timeout_seconds,
        )
    except subprocess.TimeoutExpired as exc:
        stdout = (exc.stdout or "") if isinstance(exc.stdout, str) else ""
        stderr = (exc.stderr or "") if isinstance(exc.stderr, str) else ""
        return {
            "ok": False,
            "timeout": True,
            "command": command,
            "cwd": str(workdir),
            "stdout": stdout[-MAX_COMMAND_OUTPUT:],
            "stderr": stderr[-MAX_COMMAND_OUTPUT:],
            "error": f"命令超时: {timeout_seconds}s",
        }
    except Exception as exc:
        return {"ok": False, "command": command, "cwd": str(workdir), "error": str(exc)}
    stdout = completed.stdout[-MAX_COMMAND_OUTPUT:]
    stderr = completed.stderr[-MAX_COMMAND_OUTPUT:]
    return {
        "ok": completed.returncode == 0,
        "command": command,
        "cwd": str(workdir),
        "returncode": completed.returncode,
        "stdout": stdout,
        "stderr": stderr,
        "truncated": len(completed.stdout) > MAX_COMMAND_OUTPUT or len(completed.stderr) > MAX_COMMAND_OUTPUT,
    }


def run_project_check(
    project_path: str,
    check_type: str = "auto",
    timeout_seconds: int = 180,
    confirmed: bool = False,
    permission_mode: str = "standard",
) -> dict[str, Any]:
    root = _norm(project_path)
    if not root.exists() or not root.is_dir():
        return {"ok": False, "error": f"项目目录不存在或不是目录: {root}"}
    commands = []
    if check_type in {"auto", "npm_check"} and (root / "package.json").exists():
        commands.append("npm run check")
    if check_type in {"auto", "npm_build"} and (root / "package.json").exists():
        commands.append("npm run build")
    if check_type in {"auto", "python_tests"} and (root / "sidecar").exists():
        commands.append("python -m compileall -q sidecar scripts")
    if not commands:
        return {"ok": False, "error": "没有识别到可运行的项目检查命令", "path": str(root)}
    if permission_mode != "autonomous" and not confirmed:
        return {
            "ok": False,
            "needs_confirmation": True,
            "project_check": True,
            "path": str(root),
            "check_type": check_type,
            "commands": commands,
            "message": "\n".join(
                [
                    "[CONFIRM_PROJECT_CHECK_REQUIRED]",
                    f"项目: `{root}`",
                    f"类型: {check_type}",
                    "命令:",
                    *[f"- `{command}`" for command in commands],
                    "提示: 标准模式下运行项目检查前需要确认。",
                    "[/CONFIRM_PROJECT_CHECK_REQUIRED]",
                ]
            ),
        }
    results = []
    for command in commands:
        result = run_command(command, str(root), "powershell", timeout_seconds, confirmed, permission_mode)
        results.append(result)
        if result.get("needs_confirmation"):
            return result
        if not result.get("ok"):
            break
    return {
        "ok": all(item.get("ok") for item in results),
        "path": str(root),
        "check_type": check_type,
        "results": results,
    }


def delete_path(
    target_path: str,
    target_type: str = "auto",
    recursive: bool = False,
    confirmed: bool = False,
    permission_mode: str = "standard",
) -> dict[str, Any]:
    path = _norm(target_path)
    if not path.exists():
        return {"ok": False, "error": f"路径不存在: {path}", "path": str(path)}

    actual_type = "folder" if path.is_dir() else "file"
    if target_type not in {"auto", "file", "folder"}:
        return {"ok": False, "error": "target_type must be auto, file, or folder", "path": str(path)}
    if target_type != "auto" and target_type != actual_type:
        return {
            "ok": False,
            "error": f"目标类型不匹配：工具参数是 {target_type}，实际是 {actual_type}",
            "path": str(path),
            "actual_type": actual_type,
        }
    if actual_type == "folder" and not recursive:
        return {
            "ok": False,
            "error": "拒绝删除文件夹：recursive 必须为 true，且用户必须明确要删除整个文件夹。",
            "path": str(path),
            "actual_type": actual_type,
        }

    if permission_mode != "autonomous" and not confirmed:
        target_label = "文件夹" if actual_type == "folder" else "文件"
        warning = "这会删除整个文件夹及其中所有内容。" if actual_type == "folder" else "这会删除该文件。"
        return {
            "ok": False,
            "needs_confirmation": True,
            "path": str(path),
            "actual_type": actual_type,
            "message": "\n".join([
                "[CONFIRM_DELETE_REQUIRED]",
                f"目标: `{path}`",
                f"类型: {target_label}",
                f"提示: {warning}",
                f"确认删除这个{target_label}吗？",
                "[/CONFIRM_DELETE_REQUIRED]",
            ]),
        }

    try:
        if actual_type == "folder":
            shutil.rmtree(path)
        else:
            path.unlink()
    except Exception as exc:
        return {"ok": False, "error": f"删除失败: {exc}", "path": str(path), "actual_type": actual_type}

    return {
        "ok": True,
        "path": str(path),
        "actual_type": actual_type,
        "message": f"已删除：`{path}`",
    }


def to_json(data: dict[str, Any]) -> str:
    return json.dumps(data, ensure_ascii=False, indent=2)
