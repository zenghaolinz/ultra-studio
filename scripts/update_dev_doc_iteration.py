from pathlib import Path
import sys

from docx import Document


MARKER = "8.1 v0.4.0 Agent 智能编排与预览稳定性迭代"
OLD_MARKER = "8.1 v0.3.0 Agent 智能编排与预览稳定性迭代"


def insert_before(reference, paragraph):
    reference._p.addprevious(paragraph._p)


def update_document(input_path: Path, output_path: Path) -> None:
    document = Document(str(input_path))

    for table in document.tables:
        for row in table.rows:
            if len(row.cells) >= 2 and row.cells[0].text.strip() == "版本":
                row.cells[1].text = "0.4.0"

    old_scope = "3D 工作区：支持图生 3D、文生 3D、快速预览/高质量模式、模型预览、导出和清理结果。"
    new_scope = "3D 工作区：支持图生 3D、文生 3D、三视图生成、多视角重建、快速预览/高质量模式、模型预览、导出和清理结果。"
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old_scope:
            paragraph.text = new_scope
        elif paragraph.text.strip() == OLD_MARKER:
            paragraph.text = MARKER

    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        document.save(str(output_path))
        return

    target = next(
        (paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "9. 后续优化方向"),
        None,
    )
    if target is None:
        raise RuntimeError("未找到“9. 后续优化方向”插入位置")

    additions = [
        (MARKER, "Heading 1"),
        ("迭代日期：2026-05-26。本次更新将连续生成任务从单步路由扩展为由 LLM 根据工具返回结果自行推进的依赖型工具编排，并修复模型预览线框切换造成的材质偏色问题。", None),
        ("Agent 工具编排", "Heading 2"),
        ("对于“生成一张图片，并生成三视图，然后生成模型”等请求，路由器将含有依赖关系的多步骤任务交给通用工具循环。模型必须等待上一工具返回真实图片路径后再调用下一工具。", None),
        ("执行链路：generate_image -> generate_multiview_images_from_image -> generate_3d_from_generated_multiview。", None),
        ("模型外观修改的能力边界", "Heading 2"),
        ("当前没有直接编辑 GLB 网格、拓扑或材质的工具。对于“把刚才这个模型改成金属材质，生成三视图，然后重建模型”，系统会找到模型关联的活跃源图或预览图，先由 Flux 修改图片外观，再生成三视图并通过 Hunyuan3D 重建新的模型。", None),
        ("执行链路：modify_image_with_flux -> generate_multiview_images_from_image -> generate_3d_from_generated_multiview。最终结果展示重建源图与三视图路径，避免误解为直接修改原模型。", None),
        ("预览问题修复", "Heading 2"),
        ("实体/线框模式切换时，预览组件现会克隆当前网格材质后再应用线框显示参数，避免线框配色污染实体材质并导致切回后的异常颜色。", None),
        ("验证结果", "Heading 2"),
        ("已完成 Python 语法检查、TypeScript 类型检查、生产构建，以及同步/流式接口的组合任务编排回归验证。工作目录与决赛版本目录中的相关修改保持同步。", None),
    ]

    for text, style in additions:
        paragraph = document.add_paragraph(text, style=style)
        insert_before(target, paragraph)

    document.save(str(output_path))


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: update_dev_doc_iteration.py INPUT.docx OUTPUT.docx")
    update_document(Path(sys.argv[1]), Path(sys.argv[2]))
