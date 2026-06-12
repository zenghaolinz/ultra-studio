from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "UltraStudio_开发文档.docx"
SPEECH_PATH = DOCS / "UltraStudio_5分钟演讲稿.md"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "171615") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei UI"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    r.font.size = Pt(9.5)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, "171615")
        set_cell_text(cell, header, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    document.add_paragraph()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_dev_doc() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei UI"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Microsoft YaHei UI"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
        styles[name].font.color.rgb = RGBColor.from_string("171615")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Ultra Studio 开发文档")
    run.font.name = "Microsoft YaHei UI"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    run.font.size = Pt(22)
    run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("本地 AI 3D 创作工作台 / Agent + ComfyUI + Tauri")
    r.font.name = "Microsoft YaHei UI"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("615F58")

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["版本", "0.2.0"],
            ["参赛定位", "面向 AI + 技能创作场景的本地 3D 资产创作与展示工作台"],
            ["核心链路", "需求输入 -> 概念图生成/编辑 -> 图生/文生 3D -> 预览检查 -> 导出归档 -> 展示材料"],
            ["技术栈", "Tauri 2、React、Three.js、FastAPI、SQLite、ComfyUI、Flux、Hunyuan3D、python-docx"],
        ],
    )

    add_heading(doc, "1. 项目概述")
    doc.add_paragraph(
        "Ultra Studio 是一个本地桌面 Agent 工作台，目标是把生成式 AI 能力落到真实的 3D 资产创作流程中。"
        "应用将自然语言对话、图片生成与编辑、图生 3D、文生 3D、3D 预览、文档读写、文件整理和 ComfyUI 状态管理整合在同一套界面中。"
        "相比只返回文件路径的生成工具，Ultra Studio 更强调结果可见、过程可控和资产可交付。"
    )

    add_heading(doc, "2. 当前功能范围")
    add_bullets(
        doc,
        [
            "Agent 对话：支持普通问答、任务理解、生成前回应、思考中/生成中状态提示和结构化结果卡片。",
            "图片工作区：提供独立的手动生图和图片编辑入口，可将生成或编辑后的图片继续转入 3D 工作区。",
            "3D 工作区：支持图生 3D、文生 3D、快速预览/高质量模式、模型预览、导出和清理结果。",
            "3D 预览卡：在聊天结果或工作区中展示模型，可切换实体、线框、重置视角并导出 GLB。",
            "场景模板：提供游戏角色、文创摆件、教学模型、工业概念件等比赛向入口。",
            "文档与文件工具：支持读取 PDF/DOCX/TXT、创建和编辑 DOCX、按路径整理文件夹内容、输出文件卡片。",
            "权限模式：标准模式下删除等危险操作需要确认；自主模式适合可信环境中的连续文件整理任务。",
            "ComfyUI 状态面板：展示本地服务状态、端口、进程归属、最近日志和启动/刷新入口。",
            "展示材料：维护比赛版 PPT 与作品说明素材，便于答辩展示和现场复盘。",
        ],
    )

    add_heading(doc, "3. 系统架构")
    add_table(
        doc,
        ["层级", "主要模块", "职责"],
        [
            ["React 前端", "ChatPanel、ImageStudio、ThreeDStudio、ComfyStatus", "负责对话、图片、3D、状态面板和结果卡片展示。"],
            ["Tauri 桌面层", "src-tauri/src/lib.rs、commands", "负责窗口、桌面集成、Sidecar 启动、日志输出和打包。"],
            ["Python Sidecar", "sidecar/main.py、routes/chat.py、routes/asset_3d.py", "负责 Agent 路由、工具调用、文件处理、任务状态和 API 服务。"],
            ["生成后端", "ComfyUI、Flux、Hunyuan3D 工作流", "负责图片生成/编辑、图生 3D、文生 3D 和贴图生成。"],
            ["本地数据", "SQLite、sidecar/data、ComfyUI output", "保存会话、资产路径、模型输出和可复用上下文。"],
        ],
    )

    add_heading(doc, "4. 关键工作流")
    add_heading(doc, "4.1 文生 3D", 2)
    add_bullets(
        doc,
        [
            "用户输入文字需求，例如“生成一个工业概念件模型”。",
            "Agent 先回应用户并进入生成中状态，而不是静默启动后端。",
            "系统使用 Flux 生成可追踪源图，再将源图送入 3D 工作流。",
            "生成完成后返回 3D 资产卡，包含 GLB、预览图、法线图、UV 贴图等信息。",
        ],
    )
    add_heading(doc, "4.2 图片到 3D", 2)
    add_bullets(
        doc,
        [
            "用户上传参考图或从图片工作区生成概念图。",
            "图片可直接转入 3D 工作区作为输入。",
            "工作区提供快速预览和高质量模式，保留快速验证能力。",
            "结果统一进入模型预览、导出和归档流程。",
        ],
    )
    add_heading(doc, "4.3 基于上一轮结果继续修改", 2)
    add_bullets(
        doc,
        [
            "当用户表达“把这只猫改成黑色”等延续需求时，系统优先识别为基于上一轮 Flux 源图的编辑。",
            "编辑后的图片再进入 3D 工作流，避免重新生成一个完全无关的新模型。",
            "当用户明确要求“全新的模型”时，系统走新的文生 3D 流程，不复用上一轮源图。",
        ],
    )

    add_heading(doc, "5. 文件与文档能力")
    doc.add_paragraph(
        "文件能力采用工具调用思路，而不是简单关键词分支。Agent 可以根据任务生成结构化操作请求，Sidecar 负责执行读取、创建、编辑、整理等本地文件任务。"
        "对于删除等高风险操作，标准模式会弹出二次确认；自主模式下可跳过确认以提高连续整理效率。"
    )
    add_bullets(
        doc,
        [
            "读取：支持 PDF、DOCX、TXT 等常见资料的摘要和重点提取。",
            "写入：支持创建 Word 文档、追加内容、保留原文件并可按需备份。",
            "整理：支持根据用户提供的路径读取文件夹，生成文件清单和整理建议。",
            "模糊路径：可基于桌面等常见位置给出相近候选，并允许用户自定义路径。",
        ],
    )

    add_heading(doc, "6. ComfyUI 与本地资源调度")
    doc.add_paragraph(
        "考虑到本机显存有限，Flux 与 Hunyuan3D 等模型无法长期同时驻留。系统在重工作流前尝试释放上一轮模型显存，并通过状态面板给出服务是否就绪、端口是否监听、进程是否由应用管理等信息。"
    )
    add_bullets(
        doc,
        [
            "EXE 启动时将 Sidecar 和 ComfyUI 日志写入本地日志文件，便于定位启动失败或卡住问题。",
            "ComfyUI 运行状态在主界面显示，避免用户误判为前端卡死。",
            "生成过程中展示节点进度和连接状态，WebSocket 异常时返回可理解的错误信息。",
            "关闭应用时应尽量同步关闭由应用托管的 ComfyUI 进程，避免残留进程影响下一次生成。",
        ],
    )

    add_heading(doc, "7. 前端体验设计")
    add_bullets(
        doc,
        [
            "对话、图片、3D 三类工作区相互独立，避免手动生图和建模功能挤占 Agent 对话空间。",
            "输入框底部显示权限模式和当前模型名称，标准/自主作为 Agent 执行权限，不再误放为普通设置项。",
            "Agent 回复前显示“思考中”，生成任务显示“生成中”流光状态，减少静默等待。",
            "3D 结果用可交互资产卡展示，而不是仅给出文件路径。",
            "PPT 采用比赛展示口吻，避免“本次新增”“评委看到的是”等面向开发者或用户解释式表达。",
        ],
    )

    add_heading(doc, "8. 参赛展示材料")
    add_table(
        doc,
        ["文件", "用途"],
        [
            ["docs/UltraStudio_项目展示_比赛版.pptx", "比赛答辩主 PPT，保留当前视觉系统并修正展示口吻。"],
            ["docs/ppt_assets/slide_01.png - slide_13.png", "PPT 页面渲染图，可用于预览和重新打包。"],
            ["docs/showcase/Ultra Studio_工业概念件_展示材料.md", "示例作品展示材料。"],
            ["docs/UltraStudio_5分钟演讲稿.md", "约 5 分钟现场讲稿。"],
        ],
    )

    add_heading(doc, "9. 后续优化方向")
    add_bullets(
        doc,
        [
            "资产库持久化：跨会话管理 GLB、源图、贴图和文档。",
            "任务日志面板：记录每一步执行、耗时、失败原因和可复现参数。",
            "工作流模板化：让文生 3D、图生 3D、图片编辑流程可配置。",
            "更多本地工具接入：扩展表格、演示文稿、压缩包和批量文件整理能力。",
            "模型资源策略：根据显存状态自动选择快速/高质量流程或提示用户释放资源。",
        ],
    )

    doc.save(DOCX_PATH)


def build_speech() -> None:
    speech = """# Ultra Studio 5 分钟比赛演讲稿

各位老师、评委好，我汇报的作品是 **Ultra Studio，本地 AI 3D 创作工作台**。

这个项目想解决的问题，不只是“能不能生成一个 3D 模型”，而是从一个想法到一个可展示、可修改、可导出的 3D 资产，中间有很多步骤是断开的。比如，用户需要先写提示词，再到图片生成工具里出概念图，再切到 ComfyUI 或其他工具做 3D 生成，生成之后还要去文件夹里找 GLB、预览图、法线图、UV 贴图。如果过程中报错，还要自己判断是模型没加载、显存不够，还是工作流节点失败。对于比赛现场或教学训练来说，这个流程太分散，展示起来也不直观。

所以 Ultra Studio 的定位，是把这些分散的步骤整合成一个本地桌面工作台。它的核心链路是：输入需求，生成概念图，转入 3D 建模，检查模型与贴图，最后导出 GLB 和展示材料。用户可以用自然语言告诉 Agent 想做什么，也可以在独立的图片工作区手动生图、编辑图片，再把图片送到 3D 工作区继续生成模型。

在产品体验上，我重点做了三个部分。

第一是 **Agent 对话与任务执行**。Agent 不只是聊天，它会根据用户的需求判断是普通问答、图片生成、文生 3D、图生 3D，还是文档和文件任务。生成开始前会先回应用户，生成过程中会显示“思考中”“生成中”和节点进度，避免后端已经跑起来了但界面没有反馈。生成完成后，也不是只返回一个文件路径，而是以模型卡片、图片卡片或文件卡片的形式展示结果。

第二是 **图片和 3D 两个独立工作区**。图片工作区支持手动生图和图片编辑，生成出的概念图可以直接转入 3D。3D 工作区保留快速预览和高质量模式，用户可以根据现场演示还是最终输出选择不同流程。模型生成后可以在界面里直接预览，切换实体、线框、重置视角，并导出 GLB 文件。

第三是 **本地资源调度和文档文件能力**。因为 Flux 和 3D 模型都运行在本地，显存压力比较大，所以我加入了 ComfyUI 状态提示、进程管理和日志入口。用户可以看到 ComfyUI 是否就绪、端口是否监听、最近日志是什么。另一方面，Agent 也可以读取 PDF、创建和编辑 Word 文档、整理文件夹内容。对于删除等危险操作，标准模式会二次确认；在可信环境里也可以切到自主模式提高效率。

从技术实现上，项目采用 Tauri + React 做桌面端，Python FastAPI 作为 Sidecar 负责 Agent 路由和工具调用，ComfyUI 负责 Flux 图片生成和 Hunyuan3D 建模。前端用 Three.js 展示 3D 模型，后端通过结构化工具调用连接文件系统、文档处理和生成工作流。整个设计的原则是：重模型本地运行，轻任务快速响应，长任务过程可见，危险操作可控。

如果现场演示，我会按一个完整作品流程来展示：先选择一个工业概念件或文创摆件场景，输入需求生成概念图；然后在图片工作区进行必要编辑；接着把图片转入 3D 工作区，选择快速预览或高质量模式生成模型；生成完成后在界面中旋转预览，确认模型、贴图和日志；最后导出 GLB，并生成作品说明和产物清单。

总结一下，Ultra Studio 不是一个单点的模型调用工具，而是一个面向技能创作的本地工作流。它把“想法、生成、修改、预览、导出、展示”连接起来，让 AI 生成能力真正落到可交付的 3D 作品上。我的汇报结束，谢谢各位老师。
"""
    SPEECH_PATH.write_text(speech, encoding="utf-8")


if __name__ == "__main__":
    build_dev_doc()
    build_speech()
    print(DOCX_PATH)
    print(SPEECH_PATH)
