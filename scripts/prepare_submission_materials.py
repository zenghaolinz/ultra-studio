from __future__ import annotations

import os
import copy
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "UltraStudio_提交材料_20260519"
DOCS_DIR = SUBMISSION / "文档材料"
CODE_DIR = SUBMISSION / "完整代码"
APP_DIR = SUBMISSION / "可执行程序"
FORM_SRC = Path.home() / "Downloads" / "附件1：湖南工业大学“AI+智赋技能·创享未来”大赛报名表.docx"


PROJECT_NAME = "Ultra Studio"
INTRO = (
    "Ultra Studio 是一款面向 AI+应用开发与软件技能赛道的本地桌面智能创作工作台，"
    "围绕“自然语言需求到可交付 3D 资产”的完整流程进行设计。项目采用 Tauri + React 构建前端，"
    "Python FastAPI 作为 Sidecar 后端负责任务路由、文件与文档工具调用、项目文件夹上下文管理，"
    "并接入 ComfyUI 工作流调度 Flux 图像生成、图像编辑和 Hunyuan3D 建模能力。用户既可以在聊天中"
    "用自然语言触发图片生成、图生 3D、文生 3D、文档读取、Word 写作和文件夹整理，也可以在独立的图片"
    "与 3D 工作区中手动上传参考图、生成概念图、切换快速预览/高质量模式、预览并导出 GLB 模型。"
    "项目更新后重点强化了类 Codex 的 Agent 工作方式：支持项目文件夹与普通对话分离，项目内新对话会默认围绕"
    "该文件夹读取 PDF、DOCX、TXT 等资料；支持附件解析、项目文档要求转生图/建模提示词、流式思考与生成状态恢复、"
    "标准/自主权限模式和危险操作确认。相比原始版本，Ultra Studio 不再只是单点调用 3D 工作流，而是把提示词生成、"
    "图像/模型生成、文档处理、资源调度、结果卡片展示与导出整理连接为一个可演示、可扩展、可交付的 AI 创作系统，"
    "适用于教学展示、比赛答辩、数字内容原型、文创摆件、游戏角色和工业概念件等场景。"
)


EXCLUDES = {
    ".git",
    "node_modules",
    "dist",
    "target",
    ".venv",
    "__pycache__",
    ".pytest_cache",
    ".mypy_cache",
    "logs",
    "outputs",
    "data",
    "renders",
    "UltraStudio_提交材料_20260519",
}

EXCLUDE_SUFFIXES = {
    ".pyc",
    ".pyo",
    ".log",
    ".db",
    ".sqlite",
    ".tmp",
    ".bak",
    ".pdb",
    ".wav",
    ".mp4",
}

CODE_ROOT_ITEMS = {
    "src",
    "src-tauri",
    "sidecar",
    "scripts",
    "index.html",
    "package.json",
    "package-lock.json",
    "tsconfig.json",
    "vite.config.ts",
    "start.bat",
    "start.ps1",
}


def set_cell_text(cell, text: str, bold: bool = False, size: int = 10):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_run_font(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, ascii_font="Microsoft YaHei", east_asia_font="Microsoft YaHei"):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def set_paragraph_font(paragraph, size=10.5, color=None, bold=False):
    for run in paragraph.runs:
        set_run_font(run)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Title", 22), ("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 77, 120)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    set_run_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(90, 100, 115)
    set_run_font(run)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        set_paragraph_font(p)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=9)
        shade_cell(cell, "E8EEF5")
        if widths:
            cell.width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def create_dev_docx(path: Path):
    doc = Document()
    configure_doc(doc)
    add_title(doc, "Ultra Studio 开发文档", "本地 AI 3D 创作工作台 / Agent + ComfyUI + Tauri")

    doc.add_heading("1. 项目概述", level=1)
    doc.add_paragraph(
        "Ultra Studio 是一个面向 3D 资产创作的本地桌面 Agent 工作台，目标是把自然语言、图片生成、"
        "图生 3D、文生 3D、3D 预览、文件与文档处理整合到一个可演示、可交付的桌面应用中。"
        "项目强调长任务过程可见、生成结果可直接预览、文件操作有权限边界，并能在比赛或教学场景中完成从想法到资产的闭环。"
    )
    doc.add_paragraph(INTRO)

    doc.add_heading("2. 技术架构", level=1)
    add_table(
        doc,
        ["层级", "主要技术", "职责"],
        [
            ["桌面壳", "Tauri 2 / Rust", "窗口管理、本地命令、文件对话框、Sidecar 通信、打包发布"],
            ["前端", "React / TypeScript / Three.js", "聊天工作台、项目侧边栏、图片与 3D 工作区、模型预览、结果卡片"],
            ["后端", "FastAPI / Python 3.12", "Agent 路由、工具调用、数据库、文件/文档处理、ComfyUI 状态管理"],
            ["数据层", "SQLite / aiosqlite", "会话、项目、短期记忆、生成上下文与迁移管理"],
            ["AI 工作流", "ComfyUI / Flux / Hunyuan3D", "图片生成、图片编辑、图生 3D、文生 3D、贴图与 GLB 产物生成"],
        ],
        [2.2, 3.4, 10.0],
    )

    doc.add_heading("3. 核心模块", level=1)
    add_bullets(
        doc,
        [
            "Agent 对话模块：支持普通问答、流式回答、思考中/生成中状态、工具调用和结构化结果卡片。",
            "项目工作区：普通对话与项目对话分离，项目内新对话默认围绕指定文件夹读取、生成和整理资料。",
            "附件与文档解析：支持 TXT、PDF、DOCX、代码文件等读取；可根据项目文档要求生成图片或 3D 模型。",
            "图片工作区：支持手动文生图、图片编辑、生成结果卡片展示，并可一键转入 3D 工作流。",
            "3D 工作区：支持文字到 3D、图片到 3D、双图融合到 3D，提供快速预览和高质量模式。",
            "ComfyUI 管理：检测本地服务是否就绪，展示端口、进程和日志，并在必要时尝试启动或刷新状态。",
            "权限与文件工具：标准模式下危险操作需要确认，自主模式适合可信环境中的批量文件整理。",
        ],
    )

    doc.add_heading("4. 关键流程", level=1)
    add_table(
        doc,
        ["流程", "输入", "处理", "输出"],
        [
            ["文生图", "自然语言或项目文档要求", "Agent 生成提示词并调用 Flux", "图片预览卡和文件路径"],
            ["图生 3D", "单张参考图", "去背景后进入 Hunyuan3D 工作流", "GLB 模型、预览图、贴图"],
            ["文生 3D", "物体/角色/产品描述", "先生成源图，再转入 3D 工作流", "可旋转预览的 3D 资产"],
            ["项目文档生图", "项目文件夹 + “根据文档生成”", "匹配并读取项目文档，生成忠实提示词", "符合文档要求的图片"],
            ["文档任务", "PDF/DOCX/TXT 或文件夹", "读取、摘要、写入 Word 或整理清单", "DOCX 文档或结构化回复"],
        ],
        [2.2, 3.0, 5.2, 5.2],
    )

    doc.add_heading("5. 数据库与上下文", level=1)
    doc.add_paragraph(
        "数据库采用 SQLite，包含 conversations、stm_entries、projects 等核心表。projects 表记录项目名称与根目录，"
        "conversations 通过 project_id 与项目关联。发送消息时，后端会解析当前 conversation 的项目路径，将项目文件夹作为上下文传给工具层。"
        "这使得用户在项目对话中可以直接说“根据新建文本文档生成图片”，系统会优先在当前项目文件夹里查找相关文档，而不是要求用户重复上传附件。"
    )

    doc.add_heading("6. 安全与可控性", level=1)
    add_bullets(
        doc,
        [
            "删除、移动等高风险文件操作在标准模式下会返回确认卡片，避免误删。",
            "项目上下文默认限制在用户选定文件夹内，减少误访问无关目录的风险。",
            "生成型长任务会先给出状态反馈，切换会话后也能恢复思考中/生成中的提示。",
            "附件与文档读取采用扩展名白名单和明确路径，避免把无关缓存或数据库文件交给模型。"
        ],
    )

    doc.add_heading("7. 部署与运行", level=1)
    add_bullets(
        doc,
        [
            "前端依赖通过 npm 安装，开发环境使用 npm run dev，生产构建使用 npm run build。",
            "桌面端通过 Tauri 构建，Release 主程序位于 src-tauri/target/release/ultra-studio.exe。",
            "Sidecar 后端依赖见 sidecar/requirements.txt，启动入口为 sidecar/main.py。",
            "ComfyUI 路径由 sidecar/config.ini 配置，工作流代码集中在 sidecar/tools/comfy_client.py 与相关路由中。",
        ],
    )

    doc.add_heading("8. 测试与验证", level=1)
    add_bullets(
        doc,
        [
            "TypeScript 与 Vite 构建通过 npm run build 验证。",
            "Python 后端核心路由通过 py_compile 验证语法和导入。",
            "Tauri release 构建已生成 Windows exe 主程序。",
            "针对项目文档生图场景，已验证当前项目文件夹中的 TXT 能被准确读取并转为白色可爱小狗提示词。",
        ],
    )

    doc.save(path)


def create_manual_docx(path: Path):
    doc = Document()
    configure_doc(doc)
    add_title(doc, "Ultra Studio 应用说明书", "面向比赛提交与现场演示的使用指南")

    doc.add_heading("1. 软件定位", level=1)
    doc.add_paragraph(
        "Ultra Studio 是一款本地 AI 3D 创作桌面软件。用户可以通过自然语言、图片附件或项目文件夹资料，"
        "调用图片生成、图片编辑、3D 建模、文档读写和文件整理工具，快速完成 AI 资产创作与展示材料整理。"
    )

    doc.add_heading("2. 运行准备", level=1)
    add_bullets(
        doc,
        [
            "确认已配置 ComfyUI，并在 sidecar/config.ini 中填写本地 ComfyUI 路径。",
            "启动应用后观察右上角 ComfyUI 状态，显示“已就绪”后再提交图片或 3D 生成任务。",
            "如需使用文档读写能力，确保 sidecar Python 环境已安装 requirements.txt 中依赖。",
            "首次运行建议先用短提示词测试图片生成，再测试图生 3D 或文生 3D。",
        ],
    )

    doc.add_heading("3. 主界面说明", level=1)
    add_table(
        doc,
        ["区域", "作用"],
        [
            ["左侧对话栏", "新建普通对话、添加项目文件夹、在项目内创建新对话、搜索历史会话。"],
            ["Agent Workspace", "显示用户消息、Agent 回复、图片卡片、3D 模型卡片和文件卡片。"],
            ["输入区", "输入自然语言、拖入或选择图片/PDF/DOCX/TXT 文件，点击发送按钮提交任务。"],
            ["图片工作区", "手动生成或编辑图片，并把图片加入 3D 输入。"],
            ["3D 工作区", "上传参考图或输入描述，生成、预览、导出 GLB 模型。"],
            ["设置", "配置模型供应商、默认模型、人格提示和工具行为。"],
        ],
        [4.0, 11.5],
    )

    doc.add_heading("4. 常用操作", level=1)
    doc.add_heading("4.1 根据文字生成图片", level=2)
    add_bullets(doc, ["在聊天输入区输入图片需求，例如“生成一只可爱的白色小狗图片”。", "系统会调用图片工作流，完成后返回图片预览卡。"])
    doc.add_heading("4.2 根据项目文档生成图片", level=2)
    add_bullets(doc, ["点击“添加项目文件夹”，选择包含 TXT/PDF/DOCX 的项目目录。", "在该项目的新对话中输入“根据新建文本文档的要求生成图片”。", "Agent 会先读取项目文件夹里的匹配文档，再调用生图工具。"])
    doc.add_heading("4.3 图片转 3D", level=2)
    add_bullets(doc, ["上传一张参考图片或使用刚生成的图片。", "点击“用于 3D”或进入 3D 工作区选择图片生成。", "等待工作流完成后，在模型卡中预览并导出 GLB。"])
    doc.add_heading("4.4 文生 3D", level=2)
    add_bullets(doc, ["输入清晰的主体、风格、材质和颜色描述。", "系统会先生成源图，再把源图送入 3D 工作流。", "适合游戏角色、文创摆件、教学模型和工业概念件。"])
    doc.add_heading("4.5 文档与文件任务", level=2)
    add_bullets(doc, ["可拖入 PDF/DOCX/TXT 并要求总结、提取重点或生成新 Word。", "可在项目对话中让 Agent 围绕项目文件夹整理资料。", "删除等危险操作在标准模式下需要确认。"])

    doc.add_heading("5. 注意事项", level=1)
    add_bullets(
        doc,
        [
            "图片和 3D 生成依赖本地模型，耗时与显卡、显存和 ComfyUI 工作流有关。",
            "如生成结果偏离需求，优先检查项目文档是否被正确选中，或直接在消息中写明主体、颜色和风格。",
            "当 ComfyUI 状态未就绪时，不建议连续提交大任务。",
            "电话、成员年级等报名表信息如未最终确定，可在提交前自行补填。",
        ],
    )

    doc.save(path)


def fill_registration_form(path: Path):
    doc = Document(str(FORM_SRC))
    table = doc.tables[0]
    set_cell_text(table.cell(0, 1), PROJECT_NAME, bold=True, size=10)

    members = [
        ("曾浩林", "", "生医学院", "", "", "项目负责人 / Agent 与后端统筹"),
        ("廖新波", "", "生医学院", "", "", "前端交互与界面实现"),
        ("闫海", "", "生医学院", "", "", "3D 工作流与测试验证"),
        ("李一帆", "", "生医学院", "", "", "文档能力与材料整理"),
        ("罗清", "", "电气学院", "", "", "模型流程与展示支持"),
    ]

    # The template only has three member rows. Insert two rows before the
    # teacher section so five members do not overwrite advisor fields.
    while len(table.rows) < 11:
        template_tr = table.rows[5]._tr
        new_tr = copy.deepcopy(template_tr)
        table.rows[6]._tr.addprevious(new_tr)

    for idx, member in enumerate(members, start=3):
        row = table.rows[idx].cells
        set_cell_text(row[0], "团队成员\n（最多5人）\n（可自行加行）", size=8)
        for col, value in enumerate(member, start=1):
            set_cell_text(row[col], value, size=9)

    teacher_header = table.rows[8].cells
    set_cell_text(teacher_header[0], "指导教师\n（无指导教师则此栏所有项均填“无”）", size=8)
    for col, value in enumerate(["姓名", "性别", "学院", "职称/职务", "联系电话", "是否同意指导"], start=1):
        set_cell_text(teacher_header[col], value, bold=True, size=8)

    teacher_row = table.rows[9].cells
    set_cell_text(teacher_row[0], "指导教师\n（无指导教师则此栏所有项均填“无”）", size=8)
    for i in range(1, 7):
        set_cell_text(teacher_row[i], "无", size=9)

    intro_cell = table.cell(10, 1)
    intro_cell.text = ""
    p = intro_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(INTRO)
    run.font.size = Pt(9)
    set_run_font(run)
    intro_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    doc.save(path)


def should_ignore(path: Path) -> bool:
    rel = path.relative_to(ROOT)
    if any(part in EXCLUDES for part in rel.parts):
        return True
    if path.suffix.lower() in EXCLUDE_SUFFIXES:
        return True
    return False


def copy_source_tree():
    CODE_DIR.mkdir(parents=True, exist_ok=True)
    for item in ROOT.iterdir():
        if item.name not in CODE_ROOT_ITEMS:
            continue
        if item.name in EXCLUDES:
            continue
        target = CODE_DIR / item.name
        if item.is_dir():
            shutil.copytree(item, target, ignore=lambda d, names: [name for name in names if should_ignore(Path(d) / name)])
        elif item.is_file() and not should_ignore(item):
            shutil.copy2(item, target)


def zip_source_tree(zip_path: Path):
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in CODE_DIR.rglob("*"):
            if file.is_file():
                zf.write(file, file.relative_to(CODE_DIR.parent))


def copy_executable():
    APP_DIR.mkdir(parents=True, exist_ok=True)
    candidates = [
        ROOT / "src-tauri" / "target" / "release" / "ultra-studio.exe",
        ROOT / "outputs" / "release" / "Ultra Studio_0.2.0_x64-setup.exe",
    ]
    for candidate in candidates:
        if candidate.exists():
            shutil.copy2(candidate, APP_DIR / candidate.name)


def write_manifest():
    lines = [
        "Ultra Studio 提交材料清单",
        "",
        "1. 完整代码/",
        "   - 包含前端 src、Tauri/Rust src-tauri、Python Sidecar 后端 sidecar、脚本与配置文件。",
        "   - 已排除 node_modules、.venv、target、dist、logs、数据库、缓存与历史输出。",
        "2. 完整代码.zip",
        "   - 与完整代码文件夹内容一致，便于上传。",
        "3. 可执行程序/",
        "   - ultra-studio.exe 为本次 release 构建生成的 Windows 主程序。",
        "4. 文档材料/",
        "   - UltraStudio_开发文档.pdf",
        "   - UltraStudio_应用说明书.docx",
        "   - 附件1_报名表_已填写.docx",
        "",
        "备注：报名表中联系电话按用户要求留空，便于提交前自行补填。",
    ]
    (SUBMISSION / "提交材料清单.txt").write_text("\n".join(lines), encoding="utf-8")


def main():
    resolved = SUBMISSION.resolve()
    if ROOT not in resolved.parents:
        raise RuntimeError(f"Refusing to write outside workspace: {resolved}")
    if SUBMISSION.exists():
        shutil.rmtree(SUBMISSION)

    DOCS_DIR.mkdir(parents=True)
    temp_dev_docx = DOCS_DIR / "UltraStudio_开发文档_用于转PDF.docx"
    create_dev_docx(temp_dev_docx)
    create_manual_docx(DOCS_DIR / "UltraStudio_应用说明书.docx")
    fill_registration_form(DOCS_DIR / "附件1_报名表_已填写.docx")
    copy_source_tree()
    zip_source_tree(SUBMISSION / "完整代码.zip")
    copy_executable()
    write_manifest()

    print(SUBMISSION)
    print(temp_dev_docx)


if __name__ == "__main__":
    main()
