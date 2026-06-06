from __future__ import annotations

import copy
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
FORM_SRC = Path.home() / "Downloads" / "附件1：湖南工业大学“AI+智赋技能·创享未来”大赛报名表.docx"
FORM_OUT = ROOT / "UltraStudio_提交材料_20260519" / "文档材料" / "附件1_报名表_已填写.docx"

PROJECT_NAME = "Ultra Studio"
INTRO = (
    "Ultra Studio 是一款面向 AI+应用开发与软件技能赛道的本地桌面智能创作工作台，"
    "围绕“自然语言需求到可交付 3D 资产”的完整流程进行设计。项目采用 Tauri + React 构建前端，"
    "Python FastAPI 作为 Sidecar 后端负责任务路由、文件与文档工具调用、项目文件夹上下文管理，"
    "并接入 ComfyUI 工作流调度 Flux 图像生成、图像编辑和 Hunyuan3D 建模能力。用户既可以在聊天中"
    "用自然语言触发图片生成、图生 3D、文生 3D、文档读取、Word 写作和文件夹整理，也可以在独立的图片"
    "与 3D 工作区中手动上传参考图、生成概念图、切换快速预览/高质量模式、预览并导出 GLB 模型。"
    "项目更新后重点强化了类 Codex 的 Agent 工作方式：支持项目文件夹与普通对话分离，项目内新对话会默认围绕"
    "该文件夹读取 PDF、DOCX、TXT 等资料；支持附件解析、项目文档要求转生图/建模提示词、流式思考与生成状态恢复、"
    "标准/自主权限模式和危险操作确认。相比原始版本，Ultra Studio 不再只是单点调用 3D 工作流，而是把提示词生成、"
    "图像/模型生成、文档处理、资源调度、结果卡片展示与导出整理连接为一个可演示、可扩展、可交付的 AI 创作系统，"
    "适用于教学展示、比赛答辩、数字内容原型、文创摆件、游戏角色和工业概念件等场景。"
)


def replace_text_preserve_format(cell, text: str) -> None:
    """Replace visible text while keeping the original cell/paragraph/run formatting."""
    if not cell.paragraphs:
        cell.add_paragraph()
    first = cell.paragraphs[0]
    if first.runs:
        first.runs[0].text = text
        for run in first.runs[1:]:
            run.text = ""
    else:
        first.add_run(text)
    for paragraph in cell.paragraphs[1:]:
        for run in paragraph.runs:
            run.text = ""


def main() -> None:
    doc = Document(str(FORM_SRC))
    table = doc.tables[0]

    # Insert two extra member rows by cloning the template's blank member row.
    # This keeps borders, shading, widths, paragraph styles, and font colors.
    for _ in range(2):
        blank_member_row = table.rows[5]._tr
        new_row = copy.deepcopy(blank_member_row)
        table.rows[6]._tr.addprevious(new_row)

    replace_text_preserve_format(table.cell(0, 1), PROJECT_NAME)

    members = [
        ("曾浩林", "", "生医学院", "", "", "项目负责人 / Agent 与后端统筹"),
        ("廖新波", "", "生医学院", "", "", "前端交互与界面实现"),
        ("闫海", "", "生医学院", "", "", "3D 工作流与测试验证"),
        ("李一帆", "", "生医学院", "", "", "文档能力与材料整理"),
        ("罗清", "", "电气学院", "", "", "模型流程与展示支持"),
    ]
    for row_index, member in enumerate(members, start=3):
        cells = table.rows[row_index].cells
        for col_index, value in enumerate(member, start=1):
            replace_text_preserve_format(cells[col_index], value)

    # Teacher rows have moved down by two rows after insertion.
    for cell in table.rows[9].cells[1:]:
        replace_text_preserve_format(cell, "无")

    replace_text_preserve_format(table.cell(10, 1), INTRO)

    FORM_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(FORM_OUT))
    print(FORM_OUT)


if __name__ == "__main__":
    main()
